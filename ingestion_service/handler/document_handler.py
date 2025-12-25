"""
Handler optimizado para procesamiento avanzado de documentos.
Implementa técnicas agnósticas de preprocesamiento:
- Contextual Injected Chunking
- Search Anchors
- Fact Density
- Entity Normalization

El embedding se genera del content_contextualized (con prefijo de contexto).
"""
import logging
import hashlib
import uuid
import re
import os
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
from datetime import datetime

# Imports seguros con fallbacks
import fitz  # PyMuPDF base - REQUERIDO
try:
    import pymupdf4llm  # Helper avanzado - OPCIONAL
    PYMUPDF4LLM_AVAILABLE = True
except ImportError:
    PYMUPDF4LLM_AVAILABLE = False
    
from docx import Document as DocxDocument  # REQUERIDO
import requests
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter

from common.handlers.base_handler import BaseHandler
from ..models import DocumentIngestionRequest, ChunkModel, DocumentType
from ..models.preprocessing_models import EnrichedChunk, PreprocessingResult, DocumentContext
from ..config.settings import IngestionSettings

# Import del nuevo handler agnóstico
try:
    from .preprocess_handler import AgnosticPreprocessHandler
    PREPROCESS_AVAILABLE = True
except ImportError:
    PREPROCESS_AVAILABLE = False

# Límites de tamaño para documentos (en bytes)
MAX_PDF_SIZE = 50 * 1024 * 1024  # 50MB
MAX_DOCX_SIZE = 20 * 1024 * 1024  # 20MB
MAX_TEXT_SIZE = 10 * 1024 * 1024  # 10MB


class DocumentHandler(BaseHandler):
    """
    Handler optimizado para procesamiento robusto de documentos.
    
    Implementa técnicas agnósticas de preprocesamiento:
    - content_contextualized: Para embeddings (mejora calidad vectorial)
    - search_anchors: Para BM25 + Full-Text (mejora recall)
    - fact_density: Para Score-Boosting (mejora ranking)
    - normalized_entities: Para filtrado estructurado
    
    Compatible con el sistema existente de Nooble8 Ingestion Service.
    """
    
    def __init__(self, app_settings: IngestionSettings):
        """Inicializa el handler con la configuración de la aplicación."""
        super().__init__(app_settings)
        
        # Cache de parsers por configuración (para fallback)
        self._parsers_cache = {}
        
        # Inicializar PreprocessHandler agnóstico
        self.preprocess_handler: Optional[AgnosticPreprocessHandler] = None
        self.preprocessing_enabled = getattr(app_settings, 'enable_document_preprocessing', False)
        
        if PREPROCESS_AVAILABLE and self.preprocessing_enabled:
            try:
                self.preprocess_handler = AgnosticPreprocessHandler(app_settings)
                self._logger.info(
                    "DocumentHandler initialized with AGNOSTIC preprocessing enabled",
                    extra={"model": getattr(app_settings, 'preprocessing_model', 'unknown')}
                )
            except Exception as e:
                self._logger.warning(
                    f"Failed to initialize AgnosticPreprocessHandler: {e}. "
                    "Falling back to standard processing."
                )
                self.preprocessing_enabled = False
        else:
            self._logger.info(
                "DocumentHandler initialized (preprocessing disabled or unavailable)"
            )
        
    def _get_parser(self, chunk_size: int, chunk_overlap: int) -> SentenceSplitter:
        """Obtiene o crea un parser con cache."""
        cache_key = f"{chunk_size}:{chunk_overlap}"
        if cache_key not in self._parsers_cache:
            self._parsers_cache[cache_key] = SentenceSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separator=" ",
                paragraph_separator="\n\n",
                secondary_chunking_regex="[^,.;。？！]+[,.;。？！]?"
            )
        return self._parsers_cache[cache_key]
    
    async def process_document(
        self,
        request: DocumentIngestionRequest,
        document_id: str,
        tenant_id: str,
        collection_id: str,
        agent_ids: List[str]
    ) -> List[ChunkModel]:
        """
        Procesa documento y retorna chunks con enriquecimiento agnóstico.
        
        Flujo con preprocessing habilitado:
        1. Extrae texto crudo del documento
        2. Divide en chunks con SentenceSplitter
        3. Genera contexto del documento (UNA VEZ)
        4. Enriquece cada chunk con el LLM
        5. Retorna ChunkModels con campos agnósticos
        
        Si está deshabilitado o falla:
        - Usa el flujo tradicional de chunking sin enriquecimiento
        
        Args:
            request: Request de ingestion con el documento
            document_id: ID del documento (UUID string)
            tenant_id: ID del tenant (UUID string)
            collection_id: ID de la colección
            agent_ids: Lista de agent IDs con acceso
            
        Returns:
            Lista de ChunkModel procesados con campos agnósticos
        """
        try:
            # Validar tamaño máximo del documento antes de procesar
            self._validate_document_size(request)
            
            # Cargar documento con método mejorado
            document, extraction_info = await self._load_document_enhanced(request)
            
            # Obtener tipo de documento como string
            doc_type_value = (
                request.document_type.value 
                if hasattr(request.document_type, 'value')
                else str(request.document_type)
            )
            
            # Primero hacer chunking tradicional
            raw_chunks = await self._create_raw_chunks(
                document=document,
                extraction_info=extraction_info,
                request=request
            )
            
            self._logger.info(
                f"Created {len(raw_chunks)} raw chunks, starting enrichment...",
                extra={"document_id": document_id, "preprocessing_enabled": self.preprocessing_enabled}
            )
            
            # ================================================================
            # NUEVO: Preprocesamiento agnóstico
            # ================================================================
            chunks = []
            preprocessing_used = False
            
            if self.preprocessing_enabled and self.preprocess_handler:
                try:
                    self._logger.info(
                        f"--- [AGNOSTIC] Starting preprocessing for: {request.document_name} ---"
                    )
                    
                    # Preparar chunks para el preprocessor
                    chunks_for_preprocessing = [
                        {
                            "content": chunk["content"],
                            "chunk_id": chunk["chunk_id"],
                            "chunk_index": chunk["chunk_index"]
                        }
                        for chunk in raw_chunks
                    ]
                    
                    # Preprocesar documento completo
                    preprocessing_result = await self.preprocess_handler.preprocess_document(
                        chunks=chunks_for_preprocessing,
                        document_id=document_id,
                        document_name=request.document_name,
                        document_text=document.text
                    )
                    
                    if preprocessing_result.was_preprocessed and preprocessing_result.chunks:
                        # Crear ChunkModels desde chunks enriquecidos
                        chunks = self._create_chunks_from_enriched(
                            enriched_chunks=preprocessing_result.chunks,
                            document_id=document_id,
                            tenant_id=tenant_id,
                            collection_id=collection_id,
                            agent_ids=agent_ids,
                            request=request,
                            extraction_info=extraction_info,
                            preprocessing_result=preprocessing_result
                        )
                        preprocessing_used = True
                        
                        self._logger.info(
                            f"Document preprocessed with AGNOSTIC techniques",
                            extra={
                                "document_id": document_id,
                                "document_name": request.document_name,
                                "total_chunks": len(chunks),
                                "avg_fact_density": round(preprocessing_result.avg_fact_density, 2),
                                "total_search_anchors": preprocessing_result.total_search_anchors,
                                "total_atomic_facts": preprocessing_result.total_atomic_facts,
                                "llm_tokens": preprocessing_result.llm_usage.get("total_tokens", 0),
                                "errors": len(preprocessing_result.processing_errors)
                            }
                        )
                        
                except Exception as e:
                    self._logger.error(
                        f"Agnostic preprocessing failed, falling back to standard chunking: {e}",
                        extra={"document_name": request.document_name},
                        exc_info=True
                    )
            
            # ================================================================
            # FALLBACK: Chunking tradicional sin enriquecimiento
            # ================================================================
            if not preprocessing_used:
                chunks = self._create_chunks_traditional(
                    raw_chunks=raw_chunks,
                    document_id=document_id,
                    tenant_id=tenant_id,
                    collection_id=collection_id,
                    agent_ids=agent_ids,
                    request=request,
                    extraction_info=extraction_info
                )
            
            self._logger.info(
                f"Document processed successfully",
                extra={
                    "document_id": document_id,
                    "document_name": request.document_name,
                    "document_type": doc_type_value,
                    "total_chunks": len(chunks),
                    "preprocessing_used": preprocessing_used,
                    "extraction_method": extraction_info.get("method")
                }
            )
            
            return chunks
            
        except Exception as e:
            self._logger.error(f"Error processing document: {e}", exc_info=True)
            raise
    
    async def _create_raw_chunks(
        self,
        document: Document,
        extraction_info: Dict[str, Any],
        request: DocumentIngestionRequest
    ) -> List[Dict[str, Any]]:
        """
        Divide el documento en chunks crudos usando SentenceSplitter.
        
        Returns:
            Lista de dicts con {content, chunk_id, chunk_index}
        """
        # Limpieza avanzada de texto (excepto markdown)
        if not extraction_info.get("is_markdown", False):
            cleaned_text = self._clean_text(document.text)
            document = Document(
                text=cleaned_text,
                metadata=document.metadata,
                id_=document.id_
            )
        
        # Obtener parser con cache
        parser = self._get_parser(
            request.rag_config.chunk_size,
            request.rag_config.chunk_overlap
        )
        
        # Parsear en chunks
        nodes = parser.get_nodes_from_documents([document])
        
        raw_chunks = []
        for idx, node in enumerate(nodes):
            chunk_content = self._clean_chunk_content(node.get_content())
            raw_chunks.append({
                "content": chunk_content,
                "chunk_id": str(uuid.uuid4()),
                "chunk_index": idx,
                "start_char_idx": getattr(node, 'start_char_idx', None),
                "end_char_idx": getattr(node, 'end_char_idx', None)
            })
        
        return raw_chunks
    
    def _create_chunks_from_enriched(
        self,
        enriched_chunks: List[EnrichedChunk],
        document_id: str,
        tenant_id: str,
        collection_id: str,
        agent_ids: List[str],
        request: DocumentIngestionRequest,
        extraction_info: Dict[str, Any],
        preprocessing_result: PreprocessingResult
    ) -> List[ChunkModel]:
        """
        Crea ChunkModels desde chunks enriquecidos por el LLM.
        
        CRÍTICO: El campo 'content' contiene content_contextualized
        (con prefijo de contexto) para que el embedding sea de mejor calidad.
        """
        doc_type_value = (
            request.document_type.value 
            if hasattr(request.document_type, 'value')
            else str(request.document_type)
        )
        
        self._logger.info(
            f"--- [AGNOSTIC] Creating {len(enriched_chunks)} ChunkModels from enriched data ---"
        )

        chunks = []
        for enriched in enriched_chunks:
            # IMPORTANTE: content = content_contextualized (con prefijo)
            # Esto es lo que se vectoriza y mejora la calidad del embedding
            chunk = ChunkModel(
                chunk_id=enriched.chunk_id,
                document_id=document_id,
                tenant_id=tenant_id,
                
                # Contenido contextualizado para embeddings
                content=enriched.content_contextualized,
                content_raw=enriched.content_raw,
                
                chunk_index=enriched.chunk_index,
                collection_id=collection_id,
                agent_ids=agent_ids if agent_ids else [],
                
                # ============================================
                # CAMPOS AGNÓSTICOS - Lo importante
                # ============================================
                search_anchors=enriched.search_anchors,
                atomic_facts=enriched.atomic_facts,
                fact_density=enriched.fact_density,
                document_nature=enriched.document_nature,
                normalized_entities=enriched.normalized_entities,
                
                # Legacy (menos importantes ahora)
                keywords=[],  # Reemplazados por search_anchors
                tags=[],      # Reemplazados por document_nature
                
                metadata={
                    "document_name": request.document_name,
                    "document_type": doc_type_value,
                    "extraction_method": extraction_info.get("method", "standard"),
                    "has_tables": extraction_info.get("has_tables", False),
                    "page_count": extraction_info.get("page_count", None),
                    "word_count": enriched.word_count,
                    "language": enriched.language,
                    "contextual_prefix": enriched.contextual_prefix,
                    "preprocessing_used": True,
                    "preprocessing_model": getattr(self.app_settings, 'preprocessing_model', 'unknown'),
                    **request.metadata
                }
            )
            chunks.append(chunk)

        # Log de ejemplo del primer chunk
        if chunks:
            first = chunks[0]
            self._logger.info(
                f"[AGNOSTIC] Sample chunk created",
                extra={
                    "chunk_id": first.chunk_id,
                    "fact_density": first.fact_density,
                    "search_anchors_count": len(first.search_anchors),
                    "atomic_facts_count": len(first.atomic_facts),
                    "document_nature": first.document_nature,
                    "content_length": len(first.content),
                    "has_contextual_prefix": bool(first.metadata.get("contextual_prefix"))
                }
            )

        return chunks
    
    def _create_chunks_traditional(
        self,
        raw_chunks: List[Dict[str, Any]],
        document_id: str,
        tenant_id: str,
        collection_id: str,
        agent_ids: List[str],
        request: DocumentIngestionRequest,
        extraction_info: Dict[str, Any]
    ) -> List[ChunkModel]:
        """
        Crea chunks usando el método tradicional sin enriquecimiento LLM.
        Usado como fallback.
        """
        doc_type_value = (
            request.document_type.value 
            if hasattr(request.document_type, 'value')
            else str(request.document_type)
        )
        
        chunks = []
        for raw_chunk in raw_chunks:
            chunk = ChunkModel(
                chunk_id=raw_chunk["chunk_id"],
                document_id=document_id,
                tenant_id=tenant_id,
                content=raw_chunk["content"],
                content_raw=raw_chunk["content"],  # Sin contextualizar
                chunk_index=raw_chunk["chunk_index"],
                collection_id=collection_id,
                agent_ids=agent_ids if agent_ids else [],
                
                # Sin enriquecimiento agnóstico
                search_anchors=[],
                atomic_facts=[],
                fact_density=0.5,  # Valor por defecto
                document_nature="other",
                normalized_entities={},
                
                keywords=[],
                tags=[],
                
                metadata={
                    "document_name": request.document_name,
                    "document_type": doc_type_value,
                    "start_char_idx": raw_chunk.get("start_char_idx"),
                    "end_char_idx": raw_chunk.get("end_char_idx"),
                    "extraction_method": extraction_info.get("method", "standard"),
                    "has_tables": extraction_info.get("has_tables", False),
                    "page_count": extraction_info.get("page_count", None),
                    "word_count": len(raw_chunk["content"].split()),
                    "preprocessing_used": False,
                    **request.metadata
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    # =========================================================================
    # MÉTODOS DE EXTRACCIÓN (sin cambios significativos)
    # =========================================================================
    
    def _validate_document_size(self, request: DocumentIngestionRequest):
        """Valida el tamaño del documento antes de procesarlo."""
        if request.file_path:
            file_path = Path(request.file_path)
            if not file_path.exists():
                return
                
            file_size = os.path.getsize(file_path)
            
            if request.document_type == DocumentType.PDF and file_size > MAX_PDF_SIZE:
                raise ValueError(f"PDF file exceeds maximum size ({file_size} > {MAX_PDF_SIZE} bytes)")
                
            elif request.document_type == DocumentType.DOCX and file_size > MAX_DOCX_SIZE:
                raise ValueError(f"DOCX file exceeds maximum size ({file_size} > {MAX_DOCX_SIZE} bytes)")
                
            elif request.document_type in (DocumentType.TXT, DocumentType.HTML, DocumentType.MARKDOWN) and file_size > MAX_TEXT_SIZE:
                raise ValueError(f"Text file exceeds maximum size ({file_size} > {MAX_TEXT_SIZE} bytes)")
    
    async def _load_document_enhanced(
        self, 
        request: DocumentIngestionRequest
    ) -> Tuple[Document, Dict[str, Any]]:
        """
        Carga documento con sistema estratificado de extracción.
        
        Returns:
            (Document, extraction_info) con metadata sobre la extracción
        """
        content = None
        extraction_info = {
            "method": "standard", 
            "has_tables": False,
            "errors": []
        }
        
        source_value = (
            request.document_type.value 
            if hasattr(request.document_type, 'value')
            else str(request.document_type)
        )
        metadata = {
            "source": source_value,
            "document_name": request.document_name
        }
        
        try:
            if request.file_path:
                file_path = Path(request.file_path)
                if not file_path.exists():
                    raise FileNotFoundError(f"File not found: {request.file_path}")
                
                if source_value == DocumentType.PDF.value:
                    content, pdf_info = self._extract_pdf_robust(file_path)
                    metadata["pages"] = pdf_info.get("page_count", 0)
                    extraction_info.update(pdf_info)
                    
                elif source_value == DocumentType.DOCX.value:
                    content, docx_info = self._extract_docx_robust(file_path)
                    extraction_info.update(docx_info)
                    
                elif source_value == DocumentType.MARKDOWN.value:
                    content = file_path.read_text(encoding='utf-8', errors='ignore')
                    extraction_info["method"] = "markdown_native"
                    extraction_info["is_markdown"] = True
                    
                else:
                    content = file_path.read_text(encoding='utf-8', errors='ignore')
                    extraction_info["method"] = "plain_text"
                    
            elif request.url:
                response = await self._fetch_url(str(request.url))
                content = response
                metadata["url"] = str(request.url)
                extraction_info["method"] = "url_fetch"
                
            elif request.content:
                content = request.content
                extraction_info["method"] = "direct_content"
            else:
                raise ValueError("No content source provided")
            
            if not content or not content.strip():
                raise ValueError("Document is empty or contains only whitespace")
            
            metadata["extraction_info"] = extraction_info
            
            return Document(
                text=content,
                metadata=metadata,
                id_=self._generate_doc_hash(content)
            ), extraction_info
            
        except Exception as e:
            self._logger.error(f"Document loading failed: {e}")
            if request.file_path:
                try:
                    content = self._fallback_text_extraction(Path(request.file_path))
                    if content:
                        extraction_info["method"] = "emergency_fallback"
                        extraction_info["errors"].append(str(e))
                        return Document(
                            text=content,
                            metadata=metadata,
                            id_=self._generate_doc_hash(content)
                        ), extraction_info
                except:
                    pass
            raise
    
    def _extract_pdf_robust(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extrae texto de PDF con sistema estratificado de 3 niveles."""
        extraction_info = {
            "method": "unknown",
            "has_tables": False,
            "page_count": 0,
            "errors": []
        }
        
        if PYMUPDF4LLM_AVAILABLE:
            try:
                markdown_text = pymupdf4llm.to_markdown(
                    str(file_path),
                    page_chunks=False,
                    write_images=False,
                    show_progress=False
                )
                
                if markdown_text.strip():
                    with fitz.open(str(file_path)) as pdf:
                        page_count = len(pdf)
                    
                    return markdown_text, {
                        "method": "pymupdf4llm_markdown",
                        "is_markdown": True,
                        "has_tables": '|' in markdown_text,
                        "page_count": page_count
                    }
            except Exception as e:
                extraction_info["errors"].append(f"pymupdf4llm: {str(e)}")
                self._logger.warning(f"pymupdf4llm extraction failed, trying standard: {e}")
        
        try:
            return self._extract_with_pymupdf(file_path)
        except Exception as e:
            extraction_info["errors"].append(f"pymupdf_standard: {str(e)}")
            self._logger.warning(f"PyMuPDF standard extraction failed: {e}")
        
        try:
            content = self._fallback_pdf_extraction(file_path)
            with fitz.open(str(file_path)) as pdf:
                page_count = len(pdf)
            
            return content, {
                "method": "llama_index_fallback",
                "has_tables": False,
                "page_count": page_count,
                "errors": extraction_info["errors"]
            }
        except Exception as e:
            extraction_info["errors"].append(f"llama_index_fallback: {str(e)}")
            self._logger.error(f"All PDF extraction methods failed: {e}")
            raise RuntimeError(f"PDF extraction completely failed: {'; '.join(extraction_info['errors'])}")
    
    def _extract_with_pymupdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extracción mejorada con PyMuPDF estándar."""
        text_parts = []
        tables_found = False
        page_count = 0
        
        try:
            with fitz.open(str(file_path)) as pdf:
                page_count = len(pdf)
                
                for page_num, page in enumerate(pdf, 1):
                    if text_parts:
                        text_parts.append(f"\n\n--- Page {page_num} ---\n\n")
                    
                    page_text = page.get_text(sort=True)
                    if page_text.strip():
                        text_parts.append(page_text)
                    
                    try:
                        tables = page.find_tables()
                        if tables:
                            tables_found = True
                            for table in tables:
                                try:
                                    table_data = table.extract()
                                    if table_data:
                                        table_text = self._format_table_simple(table_data)
                                        if table_text:
                                            text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
                                except Exception as te:
                                    self._logger.debug(f"Table extraction error: {te}")
                    except Exception as fe:
                        self._logger.debug(f"Table finding not supported: {fe}")
            
            full_text = "".join(text_parts)
            
            if not full_text.strip():
                raise ValueError("PDF appears to be scanned or contains no extractable text")
            
            return full_text, {
                "method": "pymupdf_standard",
                "has_tables": tables_found,
                "page_count": page_count
            }
            
        except Exception as e:
            self._logger.error(f"PyMuPDF extraction error: {e}")
            raise
    
    def _format_table_simple(self, table_data: List[List]) -> str:
        """Formatea datos de tabla de manera simple pero clara."""
        if not table_data:
            return ""
        
        lines = []
        for row in table_data:
            if row and any(cell for cell in row if cell):
                clean_cells = [str(cell).strip() if cell else "" for cell in row]
                lines.append(" | ".join(clean_cells))
        
        if lines:
            return "\n".join(lines)
        return ""
    
    def _fallback_pdf_extraction(self, file_path: Path) -> str:
        """Método de respaldo usando llama_index."""
        try:
            from llama_index.core import SimpleDirectoryReader
            reader = SimpleDirectoryReader(input_files=[str(file_path)])
            docs = reader.load_data()
            return "\n\n".join([doc.text for doc in docs])
        except Exception as e:
            self._logger.error(f"llama_index fallback failed: {e}")
            raise
    
    def _extract_docx_robust(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extrae texto de DOCX con sistema de 2 niveles."""
        extraction_info = {
            "method": "unknown",
            "has_tables": False,
            "errors": []
        }
        
        try:
            return self._extract_docx_enhanced(file_path)
        except Exception as e:
            extraction_info["errors"].append(f"docx_enhanced: {str(e)}")
            self._logger.warning(f"Enhanced DOCX extraction failed, trying fallback: {e}")
        
        try:
            content = self._fallback_docx_extraction(file_path)
            return content, {
                "method": "llama_index_fallback",
                "has_tables": False,
                "errors": extraction_info["errors"]
            }
        except Exception as e:
            extraction_info["errors"].append(f"llama_index_fallback: {str(e)}")
            self._logger.error(f"All DOCX extraction methods failed: {e}")
            raise RuntimeError(f"DOCX extraction completely failed: {'; '.join(extraction_info['errors'])}")
    
    def _extract_docx_enhanced(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extracción mejorada de DOCX preservando estructura."""
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = []
            has_tables = False
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    if para.style and para.style.name and 'Heading' in para.style.name:
                        heading_level = self._get_heading_level(para.style.name)
                        paragraphs.append(f"\n{'#' * heading_level} {text}\n")
                    else:
                        paragraphs.append(text)
            
            if doc.tables:
                has_tables = True
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            row_data.append(cell_text)
                        if any(row_data):
                            table_text.append(" | ".join(row_data))
                    
                    if table_text:
                        paragraphs.append("\n[TABLE]")
                        paragraphs.append("\n".join(table_text))
                        paragraphs.append("[/TABLE]\n")
            
            content = "\n\n".join(paragraphs)
            
            if not content.strip():
                raise ValueError("DOCX file appears to be empty")
            
            return content, {
                "method": "python-docx",
                "has_tables": has_tables
            }
            
        except Exception as e:
            self._logger.error(f"DOCX extraction error: {e}")
            raise
    
    def _get_heading_level(self, style_name: str) -> int:
        """Determina el nivel de encabezado basado en el nombre del estilo."""
        if 'Heading 1' in style_name: return 1
        if 'Heading 2' in style_name: return 2
        if 'Heading 3' in style_name: return 3
        if 'Heading 4' in style_name: return 4
        if 'Heading 5' in style_name: return 5
        if 'Heading 6' in style_name: return 6
        return 3
    
    def _fallback_docx_extraction(self, file_path: Path) -> str:
        """Método de respaldo para DOCX usando llama_index."""
        try:
            from llama_index.core import SimpleDirectoryReader
            reader = SimpleDirectoryReader(input_files=[str(file_path)])
            docs = reader.load_data()
            return "\n\n".join([doc.text for doc in docs])
        except Exception as e:
            self._logger.error(f"llama_index DOCX fallback failed: {e}")
            raise
    
    def _fallback_text_extraction(self, file_path: Path) -> str:
        """Método de emergencia para cualquier tipo de archivo."""
        try:
            return file_path.read_text(encoding='utf-8', errors='ignore')
        except:
            try:
                return file_path.read_bytes().decode('utf-8', errors='replace')
            except Exception as e:
                self._logger.error(f"Emergency text extraction failed: {e}")
                return f"CONTENT EXTRACTION FAILED: {str(e)}"
    
    def _clean_text(self, text: str) -> str:
        """Limpia y normaliza el texto extraído."""
        if "[TABLE]" in text or "[/TABLE]" in text:
            return self._gentle_clean(text)
        
        text = ''.join(char for char in text if char == '\n' or char == '\t' or ord(char) >= 32)
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        lines = text.split('\n')
        lines = [line.strip() for line in lines]
        text = '\n'.join(lines)
        
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            if line and not all(c in '.-_=*~`#' for c in line):
                cleaned_lines.append(line)
            elif not line:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)

    def _gentle_clean(self, text: str) -> str:
        """Limpieza suave para preservar tablas y estructura."""
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        text = re.sub(r'    +', '  ', text)
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        return text

    def _clean_chunk_content(self, content: str) -> str:
        """Limpieza mínima específica para chunks."""
        content = content.strip()
        if '[TABLE]' not in content and '|' not in content:
            content = ' '.join(content.split())
        return content

    async def _fetch_url(self, url: str) -> str:
        """Descarga contenido desde URL."""
        try:
            response = requests.get(
                url,
                timeout=30,
                headers={'User-Agent': 'Mozilla/5.0 (compatible; Nooble8/1.0)'}
            )
            response.raise_for_status()
            return response.text
        except requests.RequestException as e:
            self._logger.error(f"Error fetching URL {url}: {e}")
            raise ValueError(f"Failed to fetch URL: {e}")
    
    def _generate_doc_hash(self, content: str) -> str:
       """Genera hash único para el documento."""
       return hashlib.sha256(content.encode()).hexdigest()[:16]
