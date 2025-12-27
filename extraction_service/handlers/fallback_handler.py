"""
Handler de fallback para extracción de documentos usando PyMuPDF.

Se usa cuando Docling falla o no está disponible.
PyMuPDF es más rápido pero menos estructurado que Docling.
"""

import logging
import time
import re
from pathlib import Path
from typing import Tuple, Optional, List
import asyncio

from common.handlers.base_handler import BaseHandler
from ..config.settings import ExtractionSettings
from ..models.extraction_models import (
    DocumentStructure,
    SectionInfo,
    TableInfo,
    ExtractionError
)

# Imports de PyMuPDF
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Imports para DOCX
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class FallbackHandler(BaseHandler):
    """
    Handler de fallback para extracción de documentos.
    
    Usa PyMuPDF para PDFs y python-docx para DOCX.
    Más rápido pero menos estructurado que Docling.
    """
    
    def __init__(self, app_settings: ExtractionSettings):
        """Inicializa el handler de fallback."""
        super().__init__(app_settings)
        self.timeout = app_settings.fallback_timeout
        
        # Constantes configurables (o basadas en settings)
        self.chars_per_page = 3000
        self.max_sections = 50
    
    @property
    def is_available(self) -> bool:
        """Verifica si al menos un extractor está disponible."""
        return PYMUPDF_AVAILABLE or DOCX_AVAILABLE
    
    async def extract_document(
        self,
        file_path: str,
        document_type: str,
        max_pages: Optional[int] = None
    ) -> Tuple[str, DocumentStructure, Optional[ExtractionError]]:
        """
        Extrae texto de un documento usando fallback.
        
        Args:
            file_path: Ruta al archivo
            document_type: Tipo de documento (pdf, docx, txt, etc.)
            max_pages: Límite de páginas a procesar
            
        Returns:
            Tuple de (text, structure, error_if_any)
        """
        path = Path(file_path)
        
        if not path.exists():
            return "", DocumentStructure(), ExtractionError(
                error_type="FileNotFoundError",
                error_message=f"File not found: {file_path}",
                stage="file_access",
                recoverable=False
            )
        
        start_time = time.time()
        
        try:
            doc_type = document_type.lower()
            
            if doc_type == "pdf":
                text, structure = await self._extract_pdf(path, max_pages)
            elif doc_type in ("docx", "doc"):
                text, structure = await self._extract_docx(path)
            elif doc_type in ("txt", "text"):
                text, structure = await self._extract_text(path)
            elif doc_type in ("html", "htm"):
                text, structure = await self._extract_html(path)
            elif doc_type in ("md", "markdown"):
                text, structure = await self._extract_markdown(path)
            else:
                # Intento genérico de texto plano
                text, structure = await self._extract_text(path)
            
            elapsed_ms = int((time.time() - start_time) * 1000)
            
            self._logger.info(
                f"Fallback extraction completed",
                extra={
                    "file_path": file_path,
                    "document_type": doc_type,
                    "word_count": structure.word_count,
                    "elapsed_ms": elapsed_ms
                }
            )
            
            return text, structure, None
            
        except Exception as e:
            self._logger.error(f"Fallback extraction error: {e}", exc_info=True)
            return "", DocumentStructure(), ExtractionError(
                error_type=type(e).__name__,
                error_message=str(e),
                stage="fallback_extraction",
                recoverable=False
            )
    
    async def _extract_pdf(
        self, 
        path: Path, 
        max_pages: Optional[int] = None
    ) -> Tuple[str, DocumentStructure]:
        """Extrae texto de PDF usando PyMuPDF."""
        if not PYMUPDF_AVAILABLE:
            raise RuntimeError("PyMuPDF not available")
        
        text_parts = []
        tables_found = []
        sections = []
        page_count = 0
        
        with fitz.open(str(path)) as pdf:
            page_count = len(pdf)
            pages_to_process = min(page_count, max_pages) if max_pages else page_count
            
            for page_num in range(pages_to_process):
                page = pdf[page_num]
                
                # Extraer texto
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(f"\n\n--- Page {page_num + 1} ---\n\n")
                    text_parts.append(page_text)
                
                # Intentar detectar tablas
                try:
                    tables = page.find_tables()
                    for i, table in enumerate(tables):
                        try:
                            data = table.extract()
                            if data:
                                table_md = self._format_table_as_markdown(data)
                                text_parts.append(f"\n\n{table_md}\n\n")
                                tables_found.append(TableInfo(
                                    table_index=len(tables_found),
                                    rows=len(data),
                                    cols=len(data[0]) if data else 0,
                                    start_char=sum(len(p) for p in text_parts),
                                    has_header=True,
                                    markdown_content=table_md
                                ))
                        except Exception:
                            pass
                except Exception:
                    pass
        
        full_text = "".join(text_parts)
        
        # Detectar secciones desde el texto
        sections = self._detect_sections_from_text(full_text)
        
        # Limpiar texto
        full_text = self._clean_text(full_text)
        
        structure = DocumentStructure(
            sections=sections,
            tables=tables_found,
            tables_count=len(tables_found),
            page_count=page_count,
            word_count=len(full_text.split()),
            char_count=len(full_text)
        )
        
        return full_text, structure
    
    async def _extract_docx(self, path: Path) -> Tuple[str, DocumentStructure]:
        """Extrae texto de DOCX usando python-docx."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available")
        
        doc = DocxDocument(str(path))
        text_parts = []
        sections = []
        tables_found = []
        char_pos = 0
        
        # Procesar párrafos
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detectar headings
            if para.style and para.style.name:
                style_name = para.style.name
                if 'Heading' in style_name:
                    level = self._get_heading_level(style_name)
                    md_heading = f"\n{'#' * level} {text}\n"
                    text_parts.append(md_heading)
                    
                    sections.append(SectionInfo(
                        title=text,
                        level=level,
                        start_char=char_pos,
                        parent_title=None
                    ))
                    
                    char_pos += len(md_heading)
                    continue
            
            text_parts.append(text + "\n\n")
            char_pos += len(text) + 2
        
        # Procesar tablas
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_data.append(row_data)
            
            if table_data:
                table_md = self._format_table_as_markdown(table_data)
                text_parts.append(f"\n\n{table_md}\n\n")
                tables_found.append(TableInfo(
                    table_index=i,
                    rows=len(table_data),
                    cols=len(table_data[0]) if table_data else 0,
                    start_char=char_pos,
                    has_header=True,
                    markdown_content=table_md
                ))
                char_pos += len(table_md) + 4
        
        full_text = "".join(text_parts)
        
        structure = DocumentStructure(
            sections=sections,
            tables=tables_found,
            tables_count=len(tables_found),
            page_count=max(1, len(full_text) // self.chars_per_page),
            word_count=len(full_text.split()),
            char_count=len(full_text)
        )
        
        return full_text, structure
    
    async def _extract_text(self, path: Path) -> Tuple[str, DocumentStructure]:
        """Extrae texto plano de forma no bloqueante."""
        loop = asyncio.get_event_loop()
        try:
            text = await loop.run_in_executor(None, lambda: path.read_text(encoding='utf-8', errors='ignore'))
        except Exception:
            text = await loop.run_in_executor(None, lambda: path.read_bytes().decode('utf-8', errors='replace'))
        
        sections = self._detect_sections_from_text(text)
        
        structure = DocumentStructure(
            sections=sections,
            page_count=max(1, len(text) // self.chars_per_page),
            word_count=len(text.split()),
            char_count=len(text)
        )
        
        return text, structure
    
    async def _extract_html(self, path: Path) -> Tuple[str, DocumentStructure]:
        """Extrae texto de HTML (básico)."""
        try:
            from html.parser import HTMLParser
            
            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []
                    self.in_script = False
                    self.in_style = False
                
                def handle_starttag(self, tag, attrs):
                    if tag in ('script', 'style'):
                        self.in_script = True
                    elif tag in ('p', 'div', 'br', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'):
                        self.text.append('\n')
                
                def handle_endtag(self, tag):
                    if tag in ('script', 'style'):
                        self.in_script = False
                
                def handle_data(self, data):
                    if not self.in_script and not self.in_style:
                        self.text.append(data)
            
            loop = asyncio.get_event_loop()
            html_content = await loop.run_in_executor(None, lambda: path.read_text(encoding='utf-8', errors='ignore'))
            parser = TextExtractor()
            parser.feed(html_content)
            text = ''.join(parser.text)
            
        except Exception:
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, lambda: path.read_text(encoding='utf-8', errors='ignore'))
        
        # Limpiar texto
        text = self._clean_text(text)
        sections = self._detect_sections_from_text(text)
        
        structure = DocumentStructure(
            sections=sections,
            page_count=max(1, len(text) // self.chars_per_page),
            word_count=len(text.split()),
            char_count=len(text)
        )
        
        return text, structure
    
    async def _extract_markdown(self, path: Path) -> Tuple[str, DocumentStructure]:
        """Extrae texto de Markdown de forma no bloqueante (preservando estructura)."""
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, lambda: path.read_text(encoding='utf-8', errors='ignore'))
        
        # Parsear secciones de Markdown
        sections = []
        lines = text.split('\n')
        char_pos = 0
        
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('#'):
                level = 0
                for char in stripped:
                    if char == '#':
                        level += 1
                    else:
                        break
                
                if 1 <= level <= 6:
                    title = stripped[level:].strip()
                    if title:
                        sections.append(SectionInfo(
                            title=title,
                            level=level,
                            start_char=char_pos,
                            parent_title=None
                        ))
            
            char_pos += len(line) + 1
        
        structure = DocumentStructure(
            sections=sections,
            page_count=max(1, len(text) // self.chars_per_page),
            word_count=len(text.split()),
            char_count=len(text),
            has_toc=any('contenido' in s.title.lower() or 'index' in s.title.lower() 
                       for s in sections[:5])
        )
        
        return text, structure
    
    def _format_table_as_markdown(self, data: List[List]) -> str:
        """Formatea datos de tabla como Markdown."""
        if not data:
            return ""
        
        lines = []
        
        # Header
        header = data[0]
        lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        
        # Rows
        for row in data[1:]:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
        
        return "\n".join(lines)
    
    def _get_heading_level(self, style_name: str) -> int:
        """Determina nivel de heading desde nombre de estilo."""
        for i in range(1, 7):
            if str(i) in style_name:
                return i
        return 2  # Default
    
    def _detect_sections_from_text(self, text: str) -> List[SectionInfo]:
        """Detecta secciones desde texto plano (heurística)."""
        sections = []
        lines = text.split('\n')
        char_pos = 0
        
        for line in lines:
            stripped = line.strip()
            
            # Detectar líneas que parecen títulos
            # (mayúsculas, cortas, seguidas de líneas vacías o nuevos párrafos)
            if stripped and len(stripped) < 100:
                # Es una línea en mayúsculas?
                if stripped.isupper() and len(stripped) > 3:
                    sections.append(SectionInfo(
                        title=stripped.title(),
                        level=2,
                        start_char=char_pos,
                        parent_title=None
                    ))
                # Tiene formato de título numerado? (1. Título, I. Título, etc.)
                elif re.match(r'^[\dIVXivx]+[\.\)]\s+\w', stripped):
                    title = re.sub(r'^[\dIVXivx]+[\.\)]\s+', '', stripped)
                    sections.append(SectionInfo(
                        title=title,
                        level=2,
                        start_char=char_pos,
                        parent_title=None
                    ))
            
            char_pos += len(line) + 1
        
        return sections[:self.max_sections]
    
    def _clean_text(self, text: str) -> str:
        """Limpia y normaliza texto."""
        # Eliminar caracteres de control
        text = ''.join(char for char in text if char == '\n' or char == '\t' or ord(char) >= 32)
        
        # Normalizar espacios
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Limpiar líneas
        lines = text.split('\n')
        lines = [line.strip() for line in lines]
        
        return '\n'.join(lines)
